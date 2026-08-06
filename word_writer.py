from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Inches


TEMPLATE_PATH = Path(__file__).parent / "Word範本.docx"


def set_cell_text(cell, text: str) -> None:
    """清除儲存格原有內容，再填入文字。"""
    cell.text = str(text or "")


def create_word_document(
    invoice_date: str,
    id_and_surname: str,
    invoice_number: str,
    amount: str,
    product_model: str,
    phone: str,
    email: str,
    student_card_image_bytes: bytes,
    student_card_back_image_bytes: bytes,
) -> BytesIO:
    """將確認後的資料填入 Word 範本，並回傳可下載的 Word 檔。"""

    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError("找不到 Word範本.docx，請確認檔案位於專案資料夾。")

    document = Document(TEMPLATE_PATH)

    if not document.tables:
        raise ValueError("Word 範本中找不到表格。")

    table = document.tables[0]

    # 第一列是標題，以下依照範本的儲存格位置填寫。
    set_cell_text(table.cell(1, 2), invoice_date)
    set_cell_text(table.cell(1, 4), id_and_surname)

    set_cell_text(table.cell(2, 2), invoice_number)
    set_cell_text(table.cell(2, 4), phone)

    set_cell_text(table.cell(3, 2), amount)
    set_cell_text(table.cell(3, 4), email)

    set_cell_text(table.cell(4, 2), product_model)

    # 產品描述按照你的需求固定留空。
    set_cell_text(table.cell(4, 4), "")

    # 將發票照片插入範本下方的照片區。
    image_cell = table.cell(6, 1)
    image_cell.text = ""

    image_paragraph = image_cell.paragraphs[0]

    # 插入學生證正面
    front_run = image_paragraph.add_run()
    front_stream = BytesIO(student_card_image_bytes)
    front_run.add_picture(
        front_stream,
        width=Inches(3.0),
    )

    # 正面與背面之間留一點空間
    image_paragraph.add_run("  ")

    # 插入學生證背面
    back_run = image_paragraph.add_run()
    back_stream = BytesIO(student_card_back_image_bytes)
    back_run.add_picture(
        back_stream,
        width=Inches(3.0),
)

    # 將完成的 Word 存入記憶體，方便 Streamlit 直接下載。
    output = BytesIO()
    document.save(output)
    output.seek(0)

    return output